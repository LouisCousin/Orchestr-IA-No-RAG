"""Tests unitaires pour le module export_engine."""

import pytest
from pathlib import Path

# B42: skip entire test module if python-docx or openpyxl are not installed
pytest.importorskip("docx", reason="python-docx not installed")
pytest.importorskip("openpyxl", reason="openpyxl not installed")
from docx import Document

from src.core.export_engine import ExportEngine, hex_to_rgb
from src.core.plan_parser import NormalizedPlan, PlanSection


@pytest.fixture
def sample_plan():
    plan = NormalizedPlan(title="Rapport de test", objective="Tester l'export DOCX")
    plan.sections = [
        PlanSection(id="1", title="Introduction", level=1),
        PlanSection(id="1.1", title="Contexte", level=2, parent_id="1"),
        PlanSection(id="2", title="Analyse", level=1),
        PlanSection(id="3", title="Conclusion", level=1),
    ]
    return plan


@pytest.fixture
def generated_sections():
    return {
        "1": "L'introduction présente le contexte général du rapport.",
        "1.1": "Le contexte détaillé de notre analyse.\n\n- Point un\n- Point deux\n- Point trois",
        "2": "L'analyse approfondie des données recueillies.\n\n**Résultats clés**\n\nLes résultats montrent une tendance positive.",
        "3": "En conclusion, les résultats sont prometteurs.",
    }


@pytest.fixture
def generated_sections_with_table():
    return {
        "1": "Voici un résumé des résultats :\n\n| Indicateur | Valeur | Évolution |\n|---|---|---|\n| Chiffre d'affaires | 1.2M€ | +15% |\n| Marge nette | 8.5% | +2.1pts |\n| Effectifs | 45 | +5 |\n\nCes résultats sont très encourageants.",
        "1.1": "Le contexte détaillé.",
        "2": "Analyse.",
        "3": "Conclusion.",
    }


@pytest.fixture
def engine():
    return ExportEngine()


class TestHexToRgb:
    def test_valid_hex(self):
        color = hex_to_rgb("#F0C441")
        assert color[0] == 240
        assert color[1] == 196
        assert color[2] == 65

    def test_hex_without_hash(self):
        color = hex_to_rgb("4E4E50")
        assert color[0] == 78
        assert color[1] == 78
        assert color[2] == 80


class TestExportDocx:
    def test_export_creates_file(self, engine, sample_plan, generated_sections, tmp_path):
        output = tmp_path / "output.docx"
        result = engine.export_docx(sample_plan, generated_sections, output, "Test")
        assert result.exists()
        assert result.suffix == ".docx"
        assert result.stat().st_size > 0

    def test_export_with_custom_styling(self, sample_plan, generated_sections, tmp_path):
        custom_engine = ExportEngine(styling={
            "primary_color": "#FF0000",
            "secondary_color": "#0000FF",
            "font_title": "Arial",
            "font_body": "Times New Roman",
            "font_size_title": 20,
            "font_size_body": 12,
        })
        output = tmp_path / "custom_output.docx"
        result = custom_engine.export_docx(sample_plan, generated_sections, output, "Custom")
        assert result.exists()

    def test_export_with_missing_sections(self, engine, sample_plan, tmp_path):
        # Seulement 2 sections sur 4 générées
        partial_sections = {"1": "Introduction", "3": "Conclusion"}
        output = tmp_path / "partial.docx"
        result = engine.export_docx(sample_plan, partial_sections, output, "Partiel")
        assert result.exists()

    def test_export_empty_sections(self, engine, sample_plan, tmp_path):
        output = tmp_path / "empty.docx"
        result = engine.export_docx(sample_plan, {}, output, "Vide")
        assert result.exists()

    def test_export_creates_parent_dirs(self, engine, sample_plan, generated_sections, tmp_path):
        output = tmp_path / "sub" / "dir" / "output.docx"
        result = engine.export_docx(sample_plan, generated_sections, output, "Test")
        assert result.exists()


class TestMarkdownTableParsing:
    """Tests pour le parsing et rendu des tableaux markdown."""

    def test_is_table_row(self):
        assert ExportEngine._is_table_row("| A | B | C |")
        assert ExportEngine._is_table_row("| Indicateur | Valeur | Évolution |")
        assert not ExportEngine._is_table_row("Texte normal")
        assert not ExportEngine._is_table_row("| seul pipe")
        assert not ExportEngine._is_table_row("")

    def test_is_separator_row(self):
        assert ExportEngine._is_separator_row("|---|---|---|")
        assert ExportEngine._is_separator_row("| --- | --- | --- |")
        assert ExportEngine._is_separator_row("|:---|:---:|---:|")
        assert not ExportEngine._is_separator_row("| A | B | C |")
        assert not ExportEngine._is_separator_row("Texte normal")

    def test_split_into_blocks_detects_table(self, engine):
        content = "Intro.\n\n| A | B |\n|---|---|\n| 1 | 2 |\n\nFin."
        blocks = engine._split_into_blocks(content)
        types = [b[0] for b in blocks]
        assert "table" in types

    def test_split_into_blocks_detects_bullet_list(self, engine):
        content = "Intro.\n\n- Point un\n- Point deux\n\nFin."
        blocks = engine._split_into_blocks(content)
        types = [b[0] for b in blocks]
        assert "bullet_list" in types

    def test_split_into_blocks_detects_numbered_list(self, engine):
        content = "Intro.\n\n1. Premier\n2. Deuxième\n\nFin."
        blocks = engine._split_into_blocks(content)
        types = [b[0] for b in blocks]
        assert "numbered_list" in types

    def test_split_into_blocks_detects_bold_heading(self, engine):
        content = "Intro.\n\n**Sous-titre**\n\nTexte."
        blocks = engine._split_into_blocks(content)
        types = [b[0] for b in blocks]
        assert "bold_heading" in types

    def test_split_into_blocks_detects_markdown_heading(self, engine):
        """CA2-4: ## Mon titre → heading_2 block."""
        content = "Intro.\n\n## Mon titre\n\nTexte."
        blocks = engine._split_into_blocks(content)
        found = [b for b in blocks if b[0] == "heading_2"]
        assert len(found) == 1
        assert found[0][1] == "Mon titre"

    def test_split_into_blocks_normal_text_no_heading(self, engine):
        """CA2-5: Texte normal ne produit pas de block heading."""
        content = "Texte normal sans aucun titre."
        blocks = engine._split_into_blocks(content)
        heading_blocks = [b for b in blocks if b[0].startswith("heading_")]
        assert len(heading_blocks) == 0

    def test_export_with_table(self, engine, sample_plan, generated_sections_with_table, tmp_path):
        output = tmp_path / "table_output.docx"
        result = engine.export_docx(sample_plan, generated_sections_with_table, output, "Test Table")
        assert result.exists()
        assert result.stat().st_size > 0

        # Vérifier que le DOCX contient un tableau
        doc = Document(str(result))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.rows) == 4  # 1 en-tête + 3 lignes de données
        assert len(table.columns) == 3
        # Vérifier le contenu de l'en-tête
        assert table.cell(0, 0).text.strip() == "Indicateur"
        assert table.cell(0, 1).text.strip() == "Valeur"
        # Vérifier une cellule de données
        assert table.cell(1, 0).text.strip() == "Chiffre d'affaires"

    def test_export_mixed_content(self, engine, tmp_path):
        """Teste un contenu mélangé : texte, tableau, liste, gras."""
        plan = NormalizedPlan(title="Test", objective="Test mixte")
        plan.sections = [PlanSection(id="1", title="Section", level=1)]

        content = (
            "Paragraphe d'introduction.\n\n"
            "**Résultats**\n\n"
            "| Nom | Score |\n"
            "|---|---|\n"
            "| Alice | 95 |\n"
            "| Bob | 87 |\n\n"
            "- Point A\n"
            "- Point B\n\n"
            "1. Étape une\n"
            "2. Étape deux\n\n"
            "Paragraphe de conclusion avec du **gras** et de l'*italique*."
        )
        sections = {"1": content}

        output = tmp_path / "mixed.docx"
        result = engine.export_docx(plan, sections, output, "Mixed")
        assert result.exists()

        doc = Document(str(result))
        assert len(doc.tables) >= 1
        # Vérifier que le tableau a le bon contenu
        table = doc.tables[0]
        assert table.cell(0, 0).text.strip() == "Nom"
        assert table.cell(1, 0).text.strip() == "Alice"

    def test_inline_bold_and_italic(self, engine, tmp_path):
        """Teste le formatage inline (gras et italique) dans les paragraphes."""
        plan = NormalizedPlan(title="Test", objective="Test inline")
        plan.sections = [PlanSection(id="1", title="Section", level=1)]
        sections = {"1": "Texte avec du **gras** et de l'*italique* dans un paragraphe."}

        output = tmp_path / "inline.docx"
        result = engine.export_docx(plan, sections, output, "Inline")
        assert result.exists()

        doc = Document(str(result))
        # Trouver le paragraphe avec le contenu (après les en-têtes)
        found = False
        for p in doc.paragraphs:
            if "gras" in p.text:
                found = True
                bold_runs = [r for r in p.runs if r.bold]
                assert any("gras" in r.text for r in bold_runs)
                break
        assert found


class TestExportMetadataExcel:
    def test_export_excel(self, engine, sample_plan, generated_sections, tmp_path):
        output = tmp_path / "metadata.xlsx"
        cost_report = {
            "total_input_tokens": 5000,
            "total_output_tokens": 2000,
            "total_cost_usd": 0.05,
            "entries": [
                {"section_id": "1", "model": "gpt-4o", "input_tokens": 1000,
                 "output_tokens": 500, "cost_usd": 0.01, "task_type": "generation"},
            ],
        }
        result = engine.export_metadata_excel(sample_plan, generated_sections, cost_report, output)
        assert result.exists()
        assert result.suffix == ".xlsx"
