"""Génération du document DOCX final avec charte graphique configurable.

Phase 2.5 : détection des marqueurs {{NEEDS_SOURCE}} résiduels avant export.
"""

import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from src.core.plan_parser import NormalizedPlan, PlanSection
from src.utils.file_utils import ensure_dir
from src.utils.reference_cleaner import clean_source_references

logger = logging.getLogger("orchestria")


# ── Phase 2.5 : Détection des marqueurs {{NEEDS_SOURCE}} ──

_NEEDS_SOURCE_PATTERN = re.compile(r'\{\{NEEDS_SOURCE:\s*(.+?)\}\}')


def detect_needs_source_markers(text: str) -> list[dict]:
    """Détecte les marqueurs {{NEEDS_SOURCE}} dans un texte.

    Args:
        text: Texte à analyser.

    Returns:
        Liste de dicts avec les clés 'full_match', 'description', 'position'.
    """
    markers = []
    for match in _NEEDS_SOURCE_PATTERN.finditer(text):
        markers.append({
            "full_match": match.group(0),
            "description": match.group(1).strip(),
            "position": match.start(),
        })
    return markers


def scan_all_sections_for_markers(generated_sections: dict) -> dict:
    """Scanne toutes les sections générées pour détecter les marqueurs.

    Args:
        generated_sections: Dict {section_id: contenu}.

    Returns:
        Dict {section_id: list[marker_dict]} pour les sections contenant des marqueurs.
    """
    results = {}
    for section_id, content in generated_sections.items():
        markers = detect_needs_source_markers(content)
        if markers:
            results[section_id] = markers
    return results


DEFAULT_STYLING = {
    "primary_color": "#F0C441",
    "secondary_color": "#4E4E50",
    "font_title": "Calibri",
    "font_body": "Calibri",
    "font_size_title": 16,
    "font_size_body": 11,
    "margin_top_cm": 2.5,
    "margin_bottom_cm": 2.5,
    "margin_left_cm": 2.5,
    "margin_right_cm": 2.5,
    "logo_path": None,
}


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convertit une couleur hex en RGBColor."""
    hex_color = hex_color.lstrip("#")
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


class ExportEngine:
    """Génère le document DOCX final."""

    def __init__(self, styling: Optional[dict] = None):
        self.styling = {**DEFAULT_STYLING, **(styling or {})}

    def export_docx(
        self,
        plan: NormalizedPlan,
        generated_sections: dict,
        output_path: Path,
        project_name: str = "",
        warn_on_markers: bool = True,
    ) -> Path:
        """Génère le document DOCX complet.

        Phase 2.5 : détecte les marqueurs {{NEEDS_SOURCE}} résiduels.
        """
        # Phase 2.5 : Vérifier les marqueurs résiduels
        if warn_on_markers:
            markers_by_section = scan_all_sections_for_markers(generated_sections)
            if markers_by_section:
                total_markers = sum(len(m) for m in markers_by_section.values())
                logger.warning(
                    f"ATTENTION : {total_markers} marqueur(s) {{{{NEEDS_SOURCE}}}} "
                    f"détecté(s) dans {len(markers_by_section)} section(s). "
                    f"Le corpus est incomplet pour ces points."
                )
                for sid, markers in markers_by_section.items():
                    for m in markers:
                        logger.warning(f"  [{sid}] {m['description']}")

        ensure_dir(output_path.parent)
        doc = Document()

        # Phase 3 : Nettoyage des [Source N] résiduels avant export
        cleaned_sections = {
            sid: clean_source_references(content)
            for sid, content in generated_sections.items()
        }

        self._setup_styles(doc)
        self._setup_margins(doc)
        self._add_cover_page(doc, plan, project_name)
        self._add_table_of_contents(doc)
        self._add_sections(doc, plan, cleaned_sections)

        doc.save(str(output_path))
        logger.info(f"Document DOCX exporté : {output_path}")
        return output_path

    def _setup_margins(self, doc: Document) -> None:
        """Configure les marges du document."""
        for section in doc.sections:
            section.top_margin = Cm(self.styling["margin_top_cm"])
            section.bottom_margin = Cm(self.styling["margin_bottom_cm"])
            section.left_margin = Cm(self.styling["margin_left_cm"])
            section.right_margin = Cm(self.styling["margin_right_cm"])

    def _setup_styles(self, doc: Document) -> None:
        """Configure les styles du document."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.styling["font_body"]
        font.size = Pt(self.styling["font_size_body"])
        font.color.rgb = hex_to_rgb(self.styling["secondary_color"])

        # Configurer les styles de titre
        for level in range(1, 4):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                heading_style = doc.styles[style_name]
                heading_font = heading_style.font
                heading_font.name = self.styling["font_title"]
                heading_font.color.rgb = hex_to_rgb(self.styling["secondary_color"])

                if level == 1:
                    heading_font.size = Pt(self.styling["font_size_title"])
                    heading_font.bold = True
                elif level == 2:
                    heading_font.size = Pt(self.styling["font_size_title"] - 2)
                    heading_font.bold = True
                elif level == 3:
                    heading_font.size = Pt(self.styling["font_size_title"] - 4)

    def _add_cover_page(self, doc: Document, plan: NormalizedPlan, project_name: str) -> None:
        """Ajoute une page de couverture."""
        # Espace vertical
        for _ in range(6):
            doc.add_paragraph("")

        # Logo (si configuré)
        logo_path = self.styling.get("logo_path")
        if logo_path and Path(logo_path).exists():
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(logo_path), width=Inches(2))
            except Exception:
                pass

        # Titre du document
        title = plan.title or project_name or "Document"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = hex_to_rgb(self.styling["primary_color"])
        run.font.name = self.styling["font_title"]

        # Sous-titre / objectif
        if plan.objective:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(plan.objective)
            run.font.size = Pt(14)
            run.font.color.rgb = hex_to_rgb(self.styling["secondary_color"])
            run.font.name = self.styling["font_body"]

        # Ligne décorative
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("─" * 40)
        run.font.color.rgb = hex_to_rgb(self.styling["primary_color"])

        # Généré par Orchestr'IA
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Généré par Orchestr'IA")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = hex_to_rgb(self.styling["secondary_color"])

        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Ajoute une table des matières."""
        doc.add_heading("Table des matières", level=1)
        doc.add_paragraph(
            "La table des matières sera mise à jour automatiquement "
            "lors de l'ouverture dans Microsoft Word (Ctrl+A puis F9)."
        )

        # Champ TOC Word
        from docx.oxml.ns import qn
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._r.append(fldChar)

        run2 = paragraph.add_run()
        instrText = run2._r.makeelement(qn("w:instrText"), {})
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._r.append(fldChar2)

        doc.add_page_break()

    def _add_sections(self, doc: Document, plan: NormalizedPlan, generated_sections: dict) -> None:
        """Ajoute les sections générées au document."""
        for section in plan.sections:
            heading_level = min(section.level, 3)  # DOCX supporte 1-9 mais on limite à 3
            doc.add_heading(f"{section.id} {section.title}", level=heading_level)

            content = generated_sections.get(section.id, "")
            if content:
                self._add_content(doc, content, section_level=heading_level)
            else:
                p = doc.add_paragraph("[Section non générée]")
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(180, 180, 180)

    def _add_content(self, doc: Document, content: str, section_level: int = 1) -> None:
        """Parse le contenu markdown et l'ajoute au document DOCX."""
        blocks = self._split_into_blocks(content)
        for block_type, block_text in blocks:
            if block_type == "table":
                self._add_table(doc, block_text)
            elif block_type.startswith("heading_"):
                # En-tête Markdown détecté (## Titre, ### Sous-titre, etc.)
                heading_level = int(block_type.split("_")[1])
                # Positionner relativement au niveau hiérarchique de la section
                effective_level = min(heading_level + section_level, 4)
                heading = doc.add_heading(block_text, level=effective_level)
                # Appliquer la charte graphique
                for run in heading.runs:
                    run.font.name = self.styling["font_title"]
                    run.font.color.rgb = hex_to_rgb(self.styling["secondary_color"])
            elif block_type == "bullet_list":
                for line in block_text.split("\n"):
                    line = line.strip().lstrip("-•").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")
            elif block_type == "numbered_list":
                for line in block_text.split("\n"):
                    line = re.sub(r"^\d+[\.\)]\s*", "", line.strip())
                    if line:
                        doc.add_paragraph(line, style="List Number")
            elif block_type == "bold_heading":
                p = doc.add_paragraph()
                run = p.add_run(block_text.strip("* ").strip())
                run.bold = True
            else:
                self._add_rich_paragraph(doc, block_text)

    def _split_into_blocks(self, content: str) -> list[tuple[str, str]]:
        """Découpe le contenu en blocs typés (paragraph, table, bullet_list, etc.)."""
        lines = content.split("\n")
        blocks: list[tuple[str, str]] = []
        current_lines: list[str] = []
        current_type = "paragraph"

        def flush():
            text = "\n".join(current_lines).strip()
            if text:
                blocks.append((current_type, text))
            current_lines.clear()

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Détection des en-têtes Markdown (# Titre, ## Sous-titre, etc.)
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                flush()
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                blocks.append((f"heading_{level}", text))
                current_type = "paragraph"
                i += 1
                continue

            # Détection d'un tableau markdown
            if self._is_table_row(stripped) and i + 1 < len(lines) and self._is_separator_row(lines[i + 1].strip()):
                flush()
                table_lines = [stripped]
                i += 1
                while i < len(lines) and (self._is_table_row(lines[i].strip()) or self._is_separator_row(lines[i].strip())):
                    table_lines.append(lines[i].strip())
                    i += 1
                blocks.append(("table", "\n".join(table_lines)))
                current_type = "paragraph"
                continue

            # Ligne vide = séparateur de blocs
            if not stripped:
                flush()
                current_type = "paragraph"
                i += 1
                continue

            # Détection de liste à puces
            if re.match(r"^[-•]\s+", stripped):
                if current_type != "bullet_list":
                    flush()
                    current_type = "bullet_list"
                current_lines.append(stripped)
                i += 1
                continue

            # Détection de liste numérotée
            if re.match(r"^\d+[\.\)]\s+", stripped):
                if current_type != "numbered_list":
                    flush()
                    current_type = "numbered_list"
                current_lines.append(stripped)
                i += 1
                continue

            # Détection de sous-titre en gras
            if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2 and len(stripped) > 4 and stripped[2:-2].strip():
                flush()
                blocks.append(("bold_heading", stripped))
                current_type = "paragraph"
                i += 1
                continue

            # Texte normal
            if current_type not in ("paragraph",):
                flush()
                current_type = "paragraph"
            current_lines.append(line)
            i += 1

        flush()
        return blocks

    @staticmethod
    def _is_table_row(line: str) -> bool:
        """Vérifie si une ligne est une ligne de tableau markdown."""
        return line.startswith("|") and line.endswith("|") and line.count("|") >= 3

    @staticmethod
    def _is_separator_row(line: str) -> bool:
        """Vérifie si une ligne est un séparateur de tableau markdown (|---|---|)."""
        return bool(re.match(r"^\|[\s\-:]+(\|[\s\-:]+)+\|$", line))

    def _add_table(self, doc: Document, table_text: str) -> None:
        """Parse un tableau markdown et l'ajoute comme tableau DOCX."""
        lines = [l.strip() for l in table_text.split("\n") if l.strip()]

        # Filtrer les lignes de séparation (|---|---|)
        data_lines = [l for l in lines if not self._is_separator_row(l)]
        if not data_lines:
            return

        # Parser les cellules
        rows = []
        for line in data_lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)

        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        # Normaliser le nombre de colonnes
        for row in rows:
            while len(row) < num_cols:
                row.append("")

        num_rows = len(rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                self._add_inline_formatting(p, cell_text)
                # En-tête en gras
                if row_idx == 0:
                    for run in p.runs:
                        run.bold = True
                # Appliquer la police du document
                for run in p.runs:
                    run.font.name = self.styling["font_body"]
                    run.font.size = Pt(self.styling["font_size_body"])

        # Ajouter un paragraphe vide après le tableau pour l'espacement
        doc.add_paragraph("")

    def _add_rich_paragraph(self, doc: Document, text: str) -> None:
        """Ajoute un paragraphe avec du formatage inline (gras, italique)."""
        p = doc.add_paragraph()
        self._add_inline_formatting(p, text)

    @staticmethod
    def _add_inline_formatting(paragraph, text: str) -> None:
        """Ajoute du texte avec formatage inline (gras **..**, italique *..*)."""
        # Pattern pour détecter **gras** et *italique*
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part:
                paragraph.add_run(part)

    def export_metadata_excel(
        self,
        plan: NormalizedPlan,
        generated_sections: dict,
        cost_report: dict,
        output_path: Path,
        rag_coverage: Optional[dict] = None,
        deferred_sections: Optional[list] = None,
    ) -> Path:
        """Exporte les métadonnées du projet en Excel (Phase 2 complet)."""
        import pandas as pd

        ensure_dir(output_path.parent)

        # Feuille : Sections (avec couverture RAG)
        sections_data = []
        for s in plan.sections:
            row = {
                "ID": s.id,
                "Titre": s.title,
                "Niveau": s.level,
                "Budget pages": s.page_budget or "",
                "Statut": s.status,
                "Longueur (car.)": len(generated_sections.get(s.id, "")),
            }
            if rag_coverage and s.id in rag_coverage:
                cov = rag_coverage[s.id]
                row["Couverture RAG"] = cov.get("level", "")
                row["Score RAG moyen"] = cov.get("avg_score", 0)
                row["Blocs pertinents"] = cov.get("num_relevant_blocks", 0)
            if deferred_sections and s.id in deferred_sections:
                row["Reportée"] = "Oui"
            sections_data.append(row)

        # Feuille : Coûts (avec fournisseur)
        costs_data = []
        for entry in cost_report.get("entries", []):
            costs_data.append({
                "Section": entry.get("section_id", ""),
                "Fournisseur": entry.get("provider", ""),
                "Modèle": entry.get("model", ""),
                "Tokens input": entry.get("input_tokens", 0),
                "Tokens output": entry.get("output_tokens", 0),
                "Coût (USD)": entry.get("cost_usd", 0),
                "Type": entry.get("task_type", ""),
            })

        # Feuille : Récapitulatif
        recap_data = [{
            "Tokens input totaux": cost_report.get("total_input_tokens", 0),
            "Tokens output totaux": cost_report.get("total_output_tokens", 0),
            "Coût total (USD)": cost_report.get("total_cost_usd", 0),
            "Coût estimé (USD)": cost_report.get("estimated_cost_usd", 0),
            "Sections générées": len(generated_sections),
            "Sections totales": len(plan.sections),
            "Sections reportées": len(deferred_sections) if deferred_sections else 0,
        }]

        # Feuille : Contenu généré (extrait par section)
        corpus_data = []
        for s in plan.sections:
            content = generated_sections.get(s.id, "")
            if content:
                corpus_data.append({
                    "Section ID": s.id,
                    "Titre": s.title,
                    "Contenu (extrait)": content[:2000] + ("..." if len(content) > 2000 else ""),
                    "Longueur totale": len(content),
                })

        with pd.ExcelWriter(str(output_path), engine="openpyxl") as writer:
            pd.DataFrame(recap_data).to_excel(writer, sheet_name="Récapitulatif", index=False)
            pd.DataFrame(sections_data).to_excel(writer, sheet_name="Sections", index=False)
            if costs_data:
                pd.DataFrame(costs_data).to_excel(writer, sheet_name="Coûts", index=False)
            if corpus_data:
                pd.DataFrame(corpus_data).to_excel(writer, sheet_name="Contenu", index=False)

        logger.info(f"Métadonnées exportées : {output_path}")
        return output_path
