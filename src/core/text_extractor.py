"""Extraction de texte multi-format avec chaîne de fallback PDF.

Supporte : PDF (docling → pymupdf → pdfplumber → PyPDF2), DOCX, HTML, TXT/MD, Excel/CSV.
Phase 2.5 : ajout de Docling comme extracteur PDF prioritaire avec structure sémantique.
Phase 4 (Perf) : parallélisation ProcessPoolExecutor + psutil pour gestion dynamique des ressources,
                  cache par hash pour éviter les ré-extractions, singleton DocumentConverter par worker.
                  compute_optimal_workers() est publique pour réutilisation par d'autres modules.
Phase 5 (Cache) : cache JSON disque dans data/cache/ indexé par MD5(contenu+mtime),
                   clear_cache() pour invalidation manuelle.
"""

import gc
import hashlib
import json
import logging
import os
from concurrent.futures import ProcessPoolExecutor, TimeoutError as FuturesTimeoutError, as_completed
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from src.utils.config import ROOT_DIR, load_default_config
from src.utils.file_utils import ensure_dir, sha256_file, sha256_text

logger = logging.getLogger("orchestria")

# Silence les warnings ScriptRunContext des workers parallèles (ProcessPoolExecutor)
# qui perdent le lien avec la session Streamlit utilisateur.
logging.getLogger("streamlit.runtime.scriptrunner.script_runner").setLevel(logging.ERROR)

CACHE_DIR = ROOT_DIR / "data" / "cache"


# --- Cache disque JSON ---


def _get_file_hash(file_path: Path) -> str:
    """Calcule un hash MD5 combinant le contenu binaire et la date de modification.

    Ce hash sert de clé de cache : si le fichier change (contenu ou mtime),
    le hash change et le cache est invalidé naturellement.
    """
    hasher = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hasher.update(chunk)
    mtime = str(file_path.stat().st_mtime)
    hasher.update(mtime.encode("utf-8"))
    return hasher.hexdigest()


def _cache_path_for(file_hash: str) -> Path:
    """Retourne le chemin du fichier cache JSON pour un hash donné."""
    return CACHE_DIR / f"{file_hash}.json"


def _save_to_cache(file_hash: str, result: "ExtractionResult") -> None:
    """Sauvegarde un ExtractionResult en JSON dans le cache disque."""
    try:
        ensure_dir(CACHE_DIR)
        data = asdict(result)
        cache_file = _cache_path_for(file_hash)
        with open(cache_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info(f"[EXTRACTION] Cache sauvegardé pour {result.source_filename} (hash={file_hash[:12]}...)")
    except Exception as e:
        logger.warning(f"Impossible de sauvegarder le cache pour {result.source_filename}: {e}")


def _load_from_cache(file_hash: str) -> Optional["ExtractionResult"]:
    """Charge un ExtractionResult depuis le cache disque JSON.

    Returns:
        L'ExtractionResult reconstitué, ou None si absent/invalide.
    """
    cache_file = _cache_path_for(file_hash)
    if not cache_file.exists():
        return None
    try:
        with open(cache_file, "r", encoding="utf-8") as f:
            data = json.load(f)
        result = ExtractionResult(**data)
        logger.info(f"[CACHE] Résultat chargé depuis le cache pour {result.source_filename} (hash={file_hash[:12]}...)")
        return result
    except Exception as e:
        logger.warning(f"Cache invalide pour hash={file_hash[:12]}..., suppression : {e}")
        cache_file.unlink(missing_ok=True)
        return None


def clear_cache(file_path: Path) -> bool:
    """Invalide le cache d'extraction pour un fichier donné.

    Supprime le fichier JSON correspondant dans data/cache/.
    Gère silencieusement le cas où le cache n'existe pas.

    Args:
        file_path: Chemin du fichier dont on veut invalider le cache.

    Returns:
        True si un fichier cache a été supprimé, False sinon.
    """
    try:
        file_hash = _get_file_hash(file_path)
        cache_file = _cache_path_for(file_hash)
        if cache_file.exists():
            cache_file.unlink()
            logger.info(f"Cache supprimé pour {file_path.name} (hash={file_hash[:12]}...)")
            return True
        logger.debug(f"Pas de cache à supprimer pour {file_path.name}")
        return False
    except Exception as e:
        logger.warning(f"Erreur lors de la suppression du cache pour {file_path.name}: {e}")
        return False


@dataclass
class ExtractionResult:
    """Résultat d'extraction de texte d'un document."""
    text: str
    page_count: int
    char_count: int
    word_count: int
    extraction_method: str
    status: str  # "success", "partial", "failed"
    source_filename: str
    source_size_bytes: int
    source_mime_type: Optional[str] = None
    source_modified: Optional[str] = None
    hash_binary: str = ""
    hash_text: str = ""
    error_message: Optional[str] = None
    metadata: dict = field(default_factory=dict)
    structure: Optional[list[dict]] = None  # Phase 2.5 : sections détectées
    # Chaque dict : {"text": str, "type": str, "page": int, "level": int}


# --- Détection des bibliothèques disponibles ---

_AVAILABLE_PDF_LIBS: list[str] = []


def _detect_pdf_libraries() -> list[str]:
    """Détecte les bibliothèques PDF disponibles, ordonnées par priorité.

    Phase 2.5 : Docling est prioritaire car il préserve la structure sémantique.
    """
    global _AVAILABLE_PDF_LIBS
    if _AVAILABLE_PDF_LIBS:
        return _AVAILABLE_PDF_LIBS

    libs = []
    try:
        from docling.document_converter import DocumentConverter  # noqa: F401
        libs.append("docling")
    except ImportError:
        pass
    try:
        import fitz  # noqa: F401
        libs.append("pymupdf")
    except ImportError:
        pass
    try:
        import pdfplumber  # noqa: F401
        libs.append("pdfplumber")
    except ImportError:
        pass
    try:
        import PyPDF2  # noqa: F401
        libs.append("PyPDF2")
    except ImportError:
        pass

    _AVAILABLE_PDF_LIBS = libs
    logger.info(f"Bibliothèques PDF disponibles : {libs if libs else 'aucune'}")
    return libs


# --- Extraction PDF par bibliothèque ---


def _detect_heading_level(label: str) -> int:
    """Détecte le niveau hiérarchique d'un élément Docling."""
    label_lower = label.lower() if label else ""
    if "title" in label_lower or "heading" in label_lower:
        # Essayer d'extraire le niveau depuis le label (ex: "section_header_1")
        for ch in reversed(label_lower):
            if ch.isdigit():
                return int(ch)
        return 1
    return 0


def _load_pdf_extraction_config() -> dict:
    """Charge la configuration pdf_extraction depuis default.yaml avec valeurs par défaut."""
    try:
        cfg = load_default_config()
        pdf_cfg = cfg.get("pdf_extraction", {}) or {}
    except Exception:
        pdf_cfg = {}
    return {
        "docling_page_batch_size": pdf_cfg.get("docling_page_batch_size", 30),
        "docling_batch_threshold": pdf_cfg.get("docling_batch_threshold", 50),
        "coverage_threshold": pdf_cfg.get("coverage_threshold", 0.80),
        "disable_page_images": pdf_cfg.get("disable_page_images", True),
        "disable_picture_classification": pdf_cfg.get("disable_picture_classification", True),
        "disable_ocr": pdf_cfg.get("disable_ocr", True),
        "docling_max_file_size_mb": pdf_cfg.get("docling_max_file_size_mb", 50),
        "docling_timeout_seconds": pdf_cfg.get("docling_timeout_seconds", 180),
    }


def _create_docling_converter(pdf_cfg: dict):
    """Crée une instance DocumentConverter avec les options de pipeline optimisées."""
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend

    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = not pdf_cfg.get("disable_ocr", True)
    pipeline_options.generate_page_images = not pdf_cfg["disable_page_images"]
    pipeline_options.generate_picture_images = not pdf_cfg.get("disable_picture_images", True)
    pipeline_options.do_picture_classification = not pdf_cfg["disable_picture_classification"]

    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(
                pipeline_options=pipeline_options,
                backend=PyPdfiumDocumentBackend,
            )
        }
    )
    return converter


def _extract_sections_from_docling_result(result) -> list[dict]:
    """Extrait les sections structurées depuis un résultat Docling."""
    doc = result.document
    sections = []
    for item in doc.iterate_items():
        text = item.text if hasattr(item, "text") else str(item)
        if not text or not text.strip():
            continue
        label = item.label if hasattr(item, "label") else "paragraph"
        page_no = None
        if hasattr(item, "prov") and item.prov:
            page_no = item.prov[0].page_no if hasattr(item.prov[0], "page_no") else None
        sections.append({
            "text": text,
            "type": str(label),
            "page": page_no,
            "level": _detect_heading_level(str(label)),
        })
    return sections


# --- Gestion dynamique des ressources (Smart Resources) ---


def compute_optimal_workers() -> int:
    """Calcule le nombre optimal de workers selon RAM et CPU disponibles.

    Formule : workers = min(CPU_COUNT, (RAM_DISPO_GB - 2GB_SECURITY) // 1.5GB_PER_WORKER)
    Fallback : toujours au moins 1 worker.

    Cette fonction est publique pour être réutilisée par d'autres modules
    (ex: corpus_acquirer pour la parallélisation de l'extraction).
    """
    try:
        import psutil
        ram_available_gb = psutil.virtual_memory().available / (1024 ** 3)
        cpu_count = os.cpu_count() or 1
        ram_workers = int((ram_available_gb - 2.0) / 1.5)
        workers = max(1, min(cpu_count, ram_workers))
        logger.info(
            f"Smart Resources : RAM dispo={ram_available_gb:.1f}GB, "
            f"CPU={cpu_count}, workers calculés={workers}"
        )
        return workers
    except ImportError:
        logger.warning("psutil non disponible, fallback à 1 worker")
        return 1
    except Exception as e:
        logger.warning(f"Erreur calcul workers : {e}, fallback à 1")
        return 1


# --- Worker Multiprocessing pour extraction Docling parallèle ---

# Variables globales du worker (initialisées une seule fois par processus)
_worker_converter = None
_worker_pdf_cfg = None


def _init_docling_worker(pdf_cfg: dict) -> None:
    """Initializer pour ProcessPoolExecutor : charge DocumentConverter une seule fois par worker."""
    global _worker_converter, _worker_pdf_cfg
    _worker_pdf_cfg = pdf_cfg
    _worker_converter = _create_docling_converter(pdf_cfg)
    logger.info("Worker Docling initialisé (DocumentConverter chargé)")


def _docling_worker_extract_batch(args: tuple) -> list[dict]:
    """Fonction worker : extrait une plage de pages depuis un PDF.

    Args:
        args: Tuple (path_str, start_page, end_page).

    Returns:
        Liste de sections extraites (dicts).
    """
    global _worker_converter
    path_str, start, end = args
    _logger = logging.getLogger("orchestria")

    try:
        from docling.datamodel.settings import PageRange
        path = Path(path_str)
        page_range: PageRange = (start, end)
        result = _worker_converter.convert(str(path), page_range=page_range)
        sections = _extract_sections_from_docling_result(result)
        _logger.info(f"  Worker lot pages {start}-{end} → {len(sections)} éléments")
        del result
        gc.collect()
        return sections
    except Exception as e:
        _logger.warning(f"  Worker échec lot pages {start}-{end}: {e}")
        return []


# --- Cache d'extraction par hash ---


def is_extraction_cached(path: Path, metadata_store) -> bool:
    """Vérifie si un fichier a déjà été extrait avec succès via son hash binaire.

    Args:
        path: Chemin du fichier à vérifier.
        metadata_store: Instance de MetadataStore.

    Returns:
        True si l'extraction est déjà complète dans le store.
    """
    if metadata_store is None:
        return False
    try:
        file_hash = sha256_file(path)
        if not file_hash:
            return False
        conn = metadata_store._get_conn()
        row = conn.execute(
            "SELECT extraction_status FROM documents WHERE hash_binary = ? AND extraction_status = 'success'",
            (file_hash,),
        ).fetchone()
        if row:
            logger.info(f"Cache hit : {path.name} déjà extrait (hash={file_hash[:12]}...)")
            return True
    except Exception as e:
        logger.debug(f"Erreur vérification cache pour {path.name}: {e}")
    return False


def _get_pdf_page_count(path: Path) -> int:
    """Obtient le nombre de pages d'un PDF via pypdfium2 ou fitz."""
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(str(path))
        count = len(pdf)
        pdf.close()
        return count
    except Exception:
        pass
    try:
        import fitz
        doc = fitz.open(str(path))
        count = len(doc)
        doc.close()
        return count
    except Exception:
        pass
    return 0


class _DoclingFileTooLarge(Exception):
    """Raised when a PDF file exceeds the configured size limit for Docling."""
    pass


class _DoclingTimeoutError(Exception):
    """Raised when Docling extraction exceeds the configured timeout."""
    pass


def _extract_pdf_docling(path: Path) -> tuple[str, int, list[dict], str | None, str, str]:
    """Extraction PDF via Docling avec structure sémantique.

    Inclut :
    - Contrôle de poids max (docling_max_file_size_mb) : bypass vers pymupdf si dépassé
    - Coupe-circuit temporel (docling_timeout_seconds) sur future.result()
    - Options de pipeline optimisées (pas d'images, backend PyPdfium2)
    - Traitement parallèle par lots via ProcessPoolExecutor pour les gros PDF (>50 pages)
    - Singleton DocumentConverter par worker (chargé une seule fois)
    - Gestion dynamique des workers via psutil (RAM/CPU)
    - Détection de couverture et rattrapage pymupdf pour les pages manquantes

    Returns:
        Tuple (texte_complet, nombre_pages, structure_sémantique, titre, méthode, statut).

    Raises:
        _DoclingFileTooLarge: si le fichier dépasse la limite de poids configurée.
    """
    pdf_cfg = _load_pdf_extraction_config()
    batch_threshold = pdf_cfg["docling_batch_threshold"]
    batch_size = pdf_cfg["docling_page_batch_size"]
    coverage_threshold = pdf_cfg["coverage_threshold"]
    max_file_size_mb = pdf_cfg["docling_max_file_size_mb"]
    timeout_seconds = pdf_cfg["docling_timeout_seconds"]

    # ── Filtre de poids : court-circuiter Docling si fichier trop lourd ──
    file_size_bytes = os.path.getsize(path)
    file_size_mb = file_size_bytes / (1024 * 1024)
    if file_size_mb > max_file_size_mb:
        logger.warning(
            f"Fichier trop lourd pour Docling : {path.name} "
            f"({file_size_mb:.1f} Mo > {max_file_size_mb} Mo), fallback pymupdf"
        )
        raise _DoclingFileTooLarge(
            f"{path.name} ({file_size_mb:.1f} Mo) dépasse la limite Docling de {max_file_size_mb} Mo"
        )

    # Déterminer le nombre total de pages
    total_pages = _get_pdf_page_count(path)
    logger.info(f"PDF {path.name} : {total_pages} pages détectées, seuil batch = {batch_threshold}")
    if total_pages == 0:
        logger.warning(f"Impossible de déterminer le nombre de pages de {path.name}, conversion en une passe")

    sections = []
    method = "docling"

    if total_pages > batch_threshold:
        # ── Mode batch parallèle : ProcessPoolExecutor avec singleton Docling ──
        max_workers = compute_optimal_workers()

        # Construire les plages de pages
        page_ranges = []
        for start in range(1, total_pages + 1, batch_size):
            end = min(start + batch_size - 1, total_pages)
            page_ranges.append((str(path), start, end))

        logger.info(
            f"PDF volumineux ({total_pages} pages), {len(page_ranges)} lots, "
            f"{max_workers} workers parallèles"
        )

        try:
            with ProcessPoolExecutor(
                max_workers=max_workers,
                initializer=_init_docling_worker,
                initargs=(pdf_cfg,),
            ) as executor:
                futures = {
                    executor.submit(_docling_worker_extract_batch, args): args
                    for args in page_ranges
                }
                for future in as_completed(futures):
                    args = futures[future]
                    try:
                        batch_sections = future.result(timeout=timeout_seconds)
                        sections.extend(batch_sections)
                    except FuturesTimeoutError:
                        logger.warning(
                            f"  Timeout Docling ({timeout_seconds}s) lot pages "
                            f"{args[1]}-{args[2]}, fallback pymupdf"
                        )
                        future.cancel()
                        raise _DoclingTimeoutError(
                            f"Timeout Docling après {timeout_seconds}s "
                            f"pour lot pages {args[1]}-{args[2]}"
                        )
                    except Exception as e:
                        logger.warning(f"  Échec lot pages {args[1]}-{args[2]}: {e}")
        except _DoclingTimeoutError:
            raise
        except Exception as e:
            logger.warning(f"Échec ProcessPoolExecutor, fallback séquentiel : {e}")
            # Fallback séquentiel en cas d'erreur du pool
            from docling.datamodel.settings import PageRange
            for path_str, start, end in page_ranges:
                try:
                    converter = _create_docling_converter(pdf_cfg)
                    page_range: PageRange = (start, end)
                    result = converter.convert(str(path), page_range=page_range)
                    batch_sections = _extract_sections_from_docling_result(result)
                    sections.extend(batch_sections)
                except Exception as batch_err:
                    logger.warning(f"  Échec lot séquentiel pages {start}-{end}: {batch_err}")
                finally:
                    try:
                        del converter
                        del result
                    except NameError:
                        pass
                    gc.collect()
    else:
        # ── Mode single-pass : PDF de taille modérée, avec timeout ──
        try:
            with ProcessPoolExecutor(
                max_workers=1,
                initializer=_init_docling_worker,
                initargs=(pdf_cfg,),
            ) as executor:
                future = executor.submit(
                    _docling_worker_extract_batch,
                    (str(path), 1, total_pages if total_pages > 0 else 9999),
                )
                try:
                    sections = future.result(timeout=timeout_seconds)
                except FuturesTimeoutError:
                    logger.warning(
                        f"Timeout Docling ({timeout_seconds}s) single-pass pour "
                        f"{path.name}, fallback pymupdf"
                    )
                    future.cancel()
                    raise _DoclingTimeoutError(
                        f"Timeout Docling après {timeout_seconds}s pour {path.name}"
                    )

            # Récupérer le page_count depuis les sections si on ne l'a pas
            if total_pages == 0 and sections:
                total_pages = max(
                    (s.get("page") or 0 for s in sections), default=1
                )
        except _DoclingTimeoutError:
            raise
        except Exception as e:
            logger.warning(f"Erreur extraction Docling single-pass pour {path.name}: {e}")
            raise

    # ── Validation basée sur le contenu (anti "faux-négatifs") ──
    # Si Docling a extrait du texte significatif (>50 chars), on accepte directement
    # sans vérifier la couverture par pages (qui peut être mal détectée).
    # Cela évite le fallback coûteux vers PyMuPDF pour les PDFs lisibles.
    preliminary_text = "\n".join(s["text"] for s in sections)
    has_content = len(preliminary_text.strip()) > 50

    if not has_content:
        # ── Pas de contenu significatif — tentative de rattrapage pymupdf ──
        if total_pages > 0:
            pages_covered = set()
            for section in sections:
                if section.get("page"):
                    pages_covered.add(section["page"])
            missing_pages = set(range(1, total_pages + 1)) - pages_covered
        else:
            missing_pages = set()

        if missing_pages:
            logger.warning(
                f"Docling n'a pas extrait de contenu significatif pour {path.name}, "
                f"rattrapage pymupdf pour {len(missing_pages)} pages"
            )
            try:
                import fitz
                doc = fitz.open(str(path))
                recovered = 0
                for page_num in sorted(missing_pages):
                    page = doc[page_num - 1]  # fitz est 0-indexed
                    page_text = page.get_text()
                    if page_text.strip():
                        sections.append({
                            "text": page_text,
                            "type": "paragraph",
                            "page": page_num,
                            "level": 0,
                        })
                        recovered += 1
                doc.close()
                if recovered:
                    logger.info(f"Rattrapage pymupdf : {recovered} pages récupérées")
                    method = "docling+pymupdf"
            except ImportError:
                logger.warning("pymupdf non disponible pour le rattrapage des pages manquantes")
            except Exception as e:
                logger.warning(f"Échec du rattrapage pymupdf : {e}")

    # ── Reconstruction du résultat ──
    sections.sort(key=lambda s: s.get("page") or 0)
    full_text = "\n".join(s["text"] for s in sections)
    page_count = total_pages if total_pages > 0 else max(
        (s.get("page") or 0 for s in sections), default=1
    )
    if page_count == 0:
        page_count = 1

    # Déterminer le statut basé sur le contenu réel
    if len(full_text.strip()) > 50:
        status = "success"
    else:
        status = "failed"

    # Extraire le titre depuis le premier heading
    title = None
    for s in sections:
        if s.get("level", 0) >= 1:
            title = s["text"].strip()
            break

    return full_text, page_count, sections, title, method, status


def _extract_pdf_metadata(path: Path) -> dict:
    """Extrait les métadonnées (auteur, date) d'un PDF via PyPDF2."""
    metadata = {}
    try:
        import PyPDF2
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            info = reader.metadata
            if info:
                if info.author:
                    metadata["author"] = info.author
                if getattr(info, "creation_date", None):
                    metadata["creation_date"] = str(info.creation_date)
    except Exception:
        pass
    return metadata


def _extract_pdf_pymupdf(path: Path) -> tuple[str, int]:
    """Extraction PDF via pymupdf (fitz)."""
    import fitz
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text())
    page_count = len(pages)
    doc.close()
    return "\n\n".join(pages), page_count


def _extract_pdf_pdfplumber(path: Path) -> tuple[str, int]:
    """Extraction PDF via pdfplumber."""
    import pdfplumber
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages), page_count


def _extract_pdf_pypdf2(path: Path) -> tuple[str, int]:
    """Extraction PDF via PyPDF2."""
    import PyPDF2
    pages = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        page_count = len(reader.pages)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages), page_count


_PDF_EXTRACTORS = {
    "pymupdf": _extract_pdf_pymupdf,
    "pdfplumber": _extract_pdf_pdfplumber,
    "PyPDF2": _extract_pdf_pypdf2,
}


def extract_pdf(path: Path) -> ExtractionResult:
    """Extrait le texte d'un PDF avec chaîne de fallback.

    Phase 2.5 : Docling est tenté en priorité pour obtenir la structure sémantique.
    """
    available = _detect_pdf_libraries()
    if not available:
        return _make_result(
            path, text="", page_count=0, method="none",
            status="failed", error="Aucune bibliothèque PDF disponible"
        )

    # Extraire les métadonnées (auteur, date) indépendamment du texte
    pdf_meta = _extract_pdf_metadata(path)

    for lib_name in available:
        try:
            if lib_name == "docling":
                # Docling retourne aussi la structure sémantique + méthode + statut
                text, page_count, structure, title, method, status = _extract_pdf_docling(path)
                if text.strip():
                    logger.info(
                        f"PDF extrait avec {method}: {path.name} "
                        f"({page_count} pages, {len(structure)} éléments, statut={status})"
                    )
                    result = _make_result(path, text=text, page_count=page_count, method=method, status=status)
                    result.structure = structure
                    if title:
                        result.metadata["title"] = title
                    # Enrichir avec les métadonnées PDF
                    result.metadata.update(pdf_meta)
                    return result
                else:
                    logger.warning(f"Extraction vide avec docling pour {path.name}, tentative suivante...")
            else:
                extractor = _PDF_EXTRACTORS[lib_name]
                text, page_count = extractor(path)
                if text.strip():
                    logger.info(f"PDF extrait avec {lib_name}: {path.name} ({page_count} pages)")
                    result = _make_result(path, text=text, page_count=page_count, method=lib_name, status="success")
                    # Enrichir avec les métadonnées PDF
                    result.metadata.update(pdf_meta)
                    return result
                else:
                    logger.warning(f"Extraction vide avec {lib_name} pour {path.name}, tentative suivante...")
        except (_DoclingFileTooLarge, _DoclingTimeoutError) as e:
            # Docling short-circuited (file too large or timeout) — skip to next lib
            logger.warning(f"Docling contourné pour {path.name}: {e}")
            continue
        except Exception as e:
            logger.warning(f"Échec extraction PDF avec {lib_name} pour {path.name}: {e}")
            continue

    return _make_result(
        path, text="", page_count=0, method="all_failed",
        status="failed", error="Toutes les bibliothèques d'extraction ont échoué"
    )


# --- Extraction DOCX ---

def extract_docx(path: Path) -> ExtractionResult:
    """Extrait le texte d'un fichier DOCX avec structure sémantique (Phase 2.5)."""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = []
        structure = []
        title = None
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
                # Détecter le type et le niveau depuis le style
                style_name = para.style.name if para.style else "Normal"
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    structure.append({
                        "text": para.text,
                        "type": "title",
                        "page": None,
                        "level": level,
                    })
                    if title is None:
                        title = para.text.strip()
                elif style_name == "Title":
                    structure.append({
                        "text": para.text,
                        "type": "title",
                        "page": None,
                        "level": 1,
                    })
                    if title is None:
                        title = para.text.strip()
                else:
                    structure.append({
                        "text": para.text,
                        "type": "paragraph",
                        "page": None,
                        "level": 0,
                    })
        text = "\n\n".join(paragraphs)
        page_count = max(1, len(text) // 3000)  # Estimation
        result = _make_result(path, text=text, page_count=page_count, method="python-docx", status="success")
        if structure:
            result.structure = structure
        if title:
            result.metadata["title"] = title
        # Extraire auteur depuis les propriétés du document DOCX
        try:
            if doc.core_properties.author:
                result.metadata["author"] = doc.core_properties.author
            if doc.core_properties.created:
                result.metadata["creation_date"] = str(doc.core_properties.created)
        except Exception:
            pass
        return result
    except Exception as e:
        return _make_result(path, text="", page_count=0, method="python-docx", status="failed", error=str(e))


# --- Extraction HTML ---

def extract_html(path_or_text: Path | str) -> ExtractionResult:
    """Extrait le texte d'un fichier HTML en supprimant les éléments non textuels."""
    try:
        from bs4 import BeautifulSoup

        if isinstance(path_or_text, Path):
            html_content = path_or_text.read_text(encoding="utf-8", errors="replace")
            source_path = path_or_text
        else:
            html_content = path_or_text
            source_path = None

        soup = BeautifulSoup(html_content, "html.parser")

        # Extract title from <title> tag or first <h1>
        title = None
        title_tag = soup.find("title")
        if title_tag and title_tag.string:
            title = title_tag.string.strip()
        if not title:
            h1_tag = soup.find("h1")
            if h1_tag:
                title = h1_tag.get_text(strip=True)

        # Extract author from <meta name="author"> tag
        author = None
        author_tag = soup.find("meta", attrs={"name": "author"})
        if author_tag and author_tag.get("content"):
            author = author_tag["content"].strip()

        for tag in soup(["script", "style", "nav", "header", "footer", "aside", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = "\n".join(lines)

        page_count = max(1, len(text) // 3000)

        if source_path:
            result = _make_result(source_path, text=text, page_count=page_count, method="beautifulsoup", status="success")
            if title:
                result.metadata["title"] = title
            if author:
                result.metadata["author"] = author
            return result
        else:
            return ExtractionResult(
                text=text, page_count=page_count,
                char_count=len(text), word_count=len(text.split()),
                extraction_method="beautifulsoup", status="success",
                source_filename="web_content", source_size_bytes=len(html_content),
                hash_binary="", hash_text=sha256_text(text) if text else "",
            )
    except Exception as e:
        if isinstance(path_or_text, Path):
            return _make_result(path_or_text, text="", page_count=0, method="beautifulsoup", status="failed", error=str(e))
        return ExtractionResult(
            text="", page_count=0, char_count=0, word_count=0,
            extraction_method="beautifulsoup", status="failed",
            source_filename="web_content", source_size_bytes=0,
            error_message=str(e),
        )


# --- Extraction Excel/CSV ---

def extract_excel(path: Path) -> ExtractionResult:
    """Extrait le texte d'un fichier Excel ou CSV."""
    try:
        import pandas as pd
        ext = path.suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(path)
            text = df.to_string(index=False)
        else:
            dfs = pd.read_excel(path, sheet_name=None)
            parts = []
            for sheet_name, df in dfs.items():
                parts.append(f"=== Feuille: {sheet_name} ===\n{df.to_string(index=False)}")
            text = "\n\n".join(parts)
        page_count = max(1, len(text) // 3000)
        return _make_result(path, text=text, page_count=page_count, method="pandas", status="success")
    except Exception as e:
        return _make_result(path, text="", page_count=0, method="pandas", status="failed", error=str(e))


# --- Extraction TXT/Markdown ---

def extract_text_file(path: Path) -> ExtractionResult:
    """Lit un fichier texte brut ou Markdown."""
    try:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
        page_count = max(1, len(text) // 3000)
        return _make_result(path, text=text, page_count=page_count, method="direct", status="success")
    except Exception as e:
        return _make_result(path, text="", page_count=0, method="direct", status="failed", error=str(e))


# --- Routeur principal ---

def extract(path: Path, metadata_store=None, force: bool = False) -> ExtractionResult:
    """Extrait le texte d'un fichier selon son extension.

    Args:
        path: Chemin du fichier à extraire.
        metadata_store: Instance optionnelle de MetadataStore pour vérifier le cache DB.
            Si le fichier est déjà extrait (hash_binary match, status=success), retourne
            un résultat "cached" sans ré-extraction.
        force: Si True, ignore le cache disque et force une nouvelle extraction.
    """
    # Vérification cache DB (MetadataStore) avant extraction
    if metadata_store is not None and is_extraction_cached(path, metadata_store):
        return _make_result(
            path, text="", page_count=0, method="cached",
            status="cached",
        )

    # Vérification cache disque JSON (data/cache/)
    if not force and path.exists():
        try:
            file_hash = _get_file_hash(path)
            cached_result = _load_from_cache(file_hash)
            if cached_result is not None:
                return cached_result
        except Exception as e:
            logger.debug(f"Erreur vérification cache disque pour {path.name}: {e}")

    # Extraction réelle
    ext = path.suffix.lower()
    if ext == ".pdf":
        result = extract_pdf(path)
    elif ext == ".docx":
        result = extract_docx(path)
    elif ext in (".html", ".htm"):
        result = extract_html(path)
    elif ext in (".xlsx", ".xls", ".csv"):
        result = extract_excel(path)
    elif ext in (".txt", ".md", ".markdown"):
        result = extract_text_file(path)
    else:
        return _make_result(
            path, text="", page_count=0, method="unsupported",
            status="failed", error=f"Format non supporté : {ext}"
        )

    # Sauvegarder en cache disque si l'extraction a réussi
    if result.status in ("success", "partial") and path.exists():
        try:
            file_hash = _get_file_hash(path)
            _save_to_cache(file_hash, result)
            logger.info(f"[EXTRACTION] {path.name} extrait avec {result.extraction_method}")
        except Exception as e:
            logger.debug(f"Impossible de mettre en cache {path.name}: {e}")

    return result


# --- Helpers ---

def _make_result(
    path: Path,
    text: str,
    page_count: int,
    method: str,
    status: str,
    error: Optional[str] = None,
) -> ExtractionResult:
    """Construit un ExtractionResult avec les métadonnées du fichier."""
    try:
        stat = path.stat()
        size = stat.st_size
        modified = datetime.fromtimestamp(stat.st_mtime).isoformat()
    except OSError:
        size = 0
        modified = None

    return ExtractionResult(
        text=text,
        page_count=page_count,
        char_count=len(text),
        word_count=len(text.split()) if text else 0,
        extraction_method=method,
        status=status,
        source_filename=path.name,
        source_size_bytes=size,
        source_modified=modified,
        hash_binary=sha256_file(path) if path.exists() and status != "failed" else "",
        hash_text=sha256_text(text) if text else "",
        error_message=error,
    )
